"""
documents.py — Remplissage des templates Word/Excel et téléchargement .docx/.xlsx
Routes réservées au super_admin (Directeur)
"""
import io
import copy
from pathlib import Path
from typing import List

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from app.core.security import require_role

router = APIRouter(prefix="/documents", tags=["Documents"])
TEMPLATES_DIR = Path(__file__).parent.parent / "templates"


# ── Helpers python-docx ───────────────────────────────────────────────────────

def _set_para_text(para, text: str):
    """Remplace le texte d'un paragraphe en conservant le style du premier run."""
    if not para.runs:
        para.add_run(text)
        return
    # Vider tous les runs sauf le premier, y mettre le texte
    for i, run in enumerate(para.runs):
        run.text = text if i == 0 else ""


def _fill_para(para, replacements: dict):
    """Remplace les placeholders dans un paragraphe run par run, puis sur le texte assemblé."""
    full = "".join(r.text for r in para.runs)
    for key, val in replacements.items():
        full = full.replace(key, str(val))
    if para.runs:
        para.runs[0].text = full
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(full)


def _fill_doc(doc, replacements: dict):
    """Applique les remplacements sur tous les paragraphes et cellules."""
    for para in doc.paragraphs:
        _fill_para(para, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _fill_para(para, replacements)


def _docx_bytes(doc) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _docx_response(doc, filename: str) -> StreamingResponse:
    return StreamingResponse(
        io.BytesIO(_docx_bytes(doc)),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


def _xlsx_response(wb, filename: str) -> StreamingResponse:
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


# ── Schémas ───────────────────────────────────────────────────────────────────

class LigneBonAchat(BaseModel):
    objet: str = ""
    fournisseur: str = ""
    prix_unitaire: float = 0
    quantite: int = 1

class BonAchatRequest(BaseModel):
    demandeur: str = ""
    service: str = ""
    eotp: str = ""
    date: str = ""
    adresse: str = ""
    lignes: List[LigneBonAchat] = []

class ConventionStageRequest(BaseModel):
    nom_etablissement: str = ""
    statut_juridique: str = ""
    siege: str = ""
    representant: str = ""
    nom_stagiaire: str = ""
    prenom_stagiaire: str = ""
    adresse_stagiaire: str = ""
    tel_stagiaire: str = ""
    email_stagiaire: str = ""
    annee_universitaire: str = ""
    diplome: str = ""
    specialite: str = ""
    date_debut: str = ""
    date_fin: str = ""
    lieu_stage: str = ""
    structure_accueil: str = ""
    encadrant: str = ""
    montant_mensuel: str = ""
    transport: str = ""
    restauration: str = ""
    imputation: str = ""

class PrestationRequest(BaseModel):
    nom: str = ""
    prenoms: str = ""
    ne_le: str = ""
    a: str = ""
    adresse: str = ""
    tel: str = ""
    emploi_fonction: str = ""
    objet_prestation: str = ""
    produits_attendus: str = ""
    date_debut: str = ""
    date_fin: str = ""
    responsable_suivi: str = ""
    montant_net: float = 0
    montant_net_lettres: str = ""
    service_ur: str = ""
    imputation_eotp: str = ""
    date_acquit: str = ""


# ── BON D'ACHAT ───────────────────────────────────────────────────────────────

@router.post("/bon-achat/docx")
def generer_bon_achat(data: BonAchatRequest, _=Depends(require_role("super_admin"))):
    from docx import Document
    from docx.oxml.ns import qn
    import copy

    tpl = TEMPLATES_DIR / "DEMANDE D'ACHAT.docx"
    if not tpl.exists():
        raise HTTPException(404, "Template introuvable")

    doc = Document(str(tpl))

    # Remplir les paragraphes d'en-tête
    replacements = {
        "NOM ET PRENOM(S) DU DEMANDEUR :": f"NOM ET PRENOM(S) DU DEMANDEUR : {data.demandeur}",
        "SERVICE ou STRUCTURE :": f"SERVICE ou STRUCTURE : {data.service}",
        "EOTP ou CENTRE DE COÛT :": f"EOTP ou CENTRE DE COÛT : {data.eotp}",
        "Dakar le :": f"Dakar le : {data.date}",
    }
    for para in doc.paragraphs:
        for key, val in replacements.items():
            if key in para.text:
                _set_para_text(para, val)

    # Remplir le tableau : ligne de données (index 1) + adresse (dernière ligne)
    if doc.tables:
        tbl = doc.tables[0]
        # Ligne de données = row index 1
        data_row = tbl.rows[1] if len(tbl.rows) > 1 else None

        if data_row and data.lignes:
            # Dupliquer la ligne de données pour chaque article supplémentaire
            from docx.oxml import OxmlElement
            ref_row_xml = data_row._tr

            for i, ligne in enumerate(data.lignes):
                if i == 0:
                    target_row = data_row
                else:
                    # Cloner la ligne de référence
                    new_tr = copy.deepcopy(ref_row_xml)
                    ref_row_xml.addnext(new_tr)
                    # Récupérer la nouvelle ligne
                    target_row = tbl.rows[i + 1]

                cells = target_row.cells
                total = ligne.prix_unitaire * ligne.quantite
                vals = [
                    ligne.objet,
                    ligne.fournisseur,
                    f"{ligne.prix_unitaire:,.0f}" if ligne.prix_unitaire else "",
                    str(ligne.quantite),
                    f"{total:,.0f}" if total else "",
                ]
                for ci, v in enumerate(vals):
                    if ci < len(cells):
                        _set_para_text(cells[ci].paragraphs[0], v)

        # Remplir adresse de livraison (dernière ligne)
        last_row = tbl.rows[-1]
        for para in last_row.cells[0].paragraphs:
            if "Adresse de Livraison" in para.text:
                _set_para_text(para, f"Adresse de Livraison : {data.adresse}")

    return _docx_response(doc, "bon-achat.docx")


# ── CONVENTION DE STAGE ───────────────────────────────────────────────────────

@router.post("/convention-stage/docx")
def generer_convention_stage(data: ConventionStageRequest, _=Depends(require_role("super_admin"))):
    from docx import Document

    tpl = TEMPLATES_DIR / "FORMULAIRE+CONVENTION+STAGE (3).docx"
    if not tpl.exists():
        raise HTTPException(404, "Template introuvable")

    doc = Document(str(tpl))
    paras = doc.paragraphs

    def set_run(para, run_idx, value):
        """Écrit la valeur dans un run précis du paragraphe."""
        if value and run_idx < len(para.runs):
            para.runs[run_idx].text = value

    def append_run(para, value):
        """Ajoute la valeur dans le dernier run."""
        if value and para.runs:
            para.runs[-1].text = value

    # ── Établissement ──────────────────────────────────────────
    # P10: "Nom de l'organisme de formation\xa0:" + run1=" " → écrire dans run1
    set_run(paras[10], 1, data.nom_etablissement)

    # P11: run0="Statut juridique\xa0", run1=": " → ajouter après run1
    if data.statut_juridique and len(paras[11].runs) >= 2:
        paras[11].runs[1].text = ": " + data.statut_juridique

    # P12: run0="Siège social\xa0: " (un seul run) → append
    if data.siege and paras[12].runs:
        paras[12].runs[0].text = paras[12].runs[0].text.rstrip() + " " + data.siege

    # P13: "Représenté par\xa0:" + run1=" " → écrire dans run1
    set_run(paras[13], 1, data.representant)

    # ── Stagiaire ───────────────────────────────────────────────
    # P17: "Nom, Prénom\xa0:" + run1=" " → écrire dans run1
    set_run(paras[17], 1, f"{data.prenom_stagiaire} {data.nom_stagiaire}")

    # P18: "Adresse\xa0:" + run1=" " → écrire dans run1
    set_run(paras[18], 1, data.adresse_stagiaire)

    # P19: run0="Tel\xa0: " (seul run) → append
    if data.tel_stagiaire and paras[19].runs:
        paras[19].runs[0].text = paras[19].runs[0].text.rstrip() + " " + data.tel_stagiaire

    # P20: run0="Email\xa0: " (seul run) → append
    if data.email_stagiaire and paras[20].runs:
        paras[20].runs[0].text = paras[20].runs[0].text.rstrip() + " " + data.email_stagiaire

    # P21: run0="Etudiant ", run1="pour l'année universitaire\xa0:", run2=" " → écrire dans run2
    set_run(paras[21], 2, data.annee_universitaire)

    # P22: run0="Diplôme préparé", run1=" ", run2=": " → ajouter après run2
    if data.diplome and len(paras[22].runs) >= 3:
        paras[22].runs[2].text = ": " + data.diplome

    # P23: run0="Spécialité", run1=" ", run2=": " → ajouter après run2
    if data.specialite and len(paras[23].runs) >= 3:
        paras[23].runs[2].text = ": " + data.specialite

    # ── Modalités ───────────────────────────────────────────────
    # P57: run1="…………………" (date début), run3="…………………." (date fin)
    if data.date_debut:
        set_run(paras[57], 1, data.date_debut)
    if data.date_fin:
        set_run(paras[57], 3, data.date_fin + ".")

    # P59: run2="……………………………." → lieu
    set_run(paras[59], 2, data.lieu_stage + ".")

    # P61: run1=" …………………………." → structure
    if data.structure_accueil and len(paras[61].runs) >= 2:
        paras[61].runs[1].text = " " + data.structure_accueil + "."

    # P64: run1="…………………………….." → encadrant scientifique
    set_run(paras[64], 1, data.encadrant + "..")

    # P65: run2="………………………" → encadrant pédagogique
    set_run(paras[65], 2, data.encadrant)

    # ── Gratification ────────────────────────────────────────────
    # P69: run1="……….." → montant mensuel
    set_run(paras[69], 1, data.montant_mensuel or "")

    # P70: run1=" …………….." → transport
    if data.transport and len(paras[70].runs) >= 2:
        paras[70].runs[1].text = " " + data.transport

    # P71: run1="……………… " → restauration
    if data.restauration and len(paras[71].runs) >= 2:
        paras[71].runs[1].text = data.restauration + " "

    # P73: run2="……………………………" → imputation
    set_run(paras[73], 2, data.imputation or "")

    # ── Signatures ──────────────────────────────────────────────
    if doc.tables:
        cells = doc.tables[0].rows[1].cells
        if cells[0].paragraphs:
            cells[0].paragraphs[0].add_run(f"\n{data.representant}")
        if cells[1].paragraphs:
            cells[1].paragraphs[0].add_run(f"\n{data.prenom_stagiaire} {data.nom_stagiaire}")

    return _docx_response(doc, "convention-stage.docx")


# ── PRESTATION DE SERVICE ─────────────────────────────────────────────────────

@router.post("/prestation-service/xlsx")
def generer_prestation_service(data: PrestationRequest, _=Depends(require_role("super_admin"))):
    import openpyxl

    tpl = TEMPLATES_DIR / "Recu de Prestation de Service.xlsx"
    if not tpl.exists():
        raise HTTPException(404, "Template introuvable")

    wb = openpyxl.load_workbook(str(tpl))
    montant = float(data.montant_net)
    impot = round(montant * 0.05, 2) if montant >= 25000 else 0
    brut = montant + impot

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        cell_map = {
            "A6": data.nom,
            "D6": data.prenoms,
            "A7": f"né(e) le : {data.ne_le}",
            "D7": f"à : {data.a}",
            "A8": f"Adresse : {data.adresse}",
            "D8": f"Tél : {data.tel}",
            "A9": f"Emploi/Fonction : {data.emploi_fonction}",
            "A10": data.objet_prestation,
            "A13": data.produits_attendus,
            "A15": f"Du : {data.date_debut}",
            "D15": f"au : {data.date_fin}",
            "A16": f"Responsable du suivi : {data.responsable_suivi}",
            "E20": montant,
            "A21": data.montant_net_lettres,
            "A25": data.service_ur,
            "A26": data.imputation_eotp,
            "A35": "POUR ACQUIT",
            "D35": data.date_acquit,
        }
        if impot > 0:
            cell_map["C23"] = impot
            cell_map["C24"] = brut

        for coord, val in cell_map.items():
            if val not in ("", 0):
                ws[coord] = val

    return _xlsx_response(wb, "prestation-service.xlsx")
